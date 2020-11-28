from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.http import HttpResponse
from django.views import View
from django.contrib import messages
from django.http import FileResponse
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import permission_required
from django.core.paginator import Paginator
import io
import os
import shutil
from django.conf import settings
import docx
import json
import datetime

import subprocess
# import models from warehousemanager app
from warehousemanager.models import *
from warehousemanager.forms import *


def index(request):
    return HttpResponse('first view')


# view displays all orders
class Orders(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_order'

    def get(self, request):
        orders_all = Order.objects.all()
        return render(request, 'warehousemanager-orders.html', locals())


class OrdersDetails(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_order'

    def get(self, request, order_id):
        order = Order.objects.get(id=order_id)
        ordered_items = order.orderitem_set.all()
        return render(request, 'warehousemanager-order-details.html', locals())


class AllOrdersDetails(LoginRequiredMixin, PermissionRequiredMixin, View):
    login_url = '/'
    permission_required = 'warehousemanager.view_order'

    def get(self, request):
        orders = Order.objects.filter(is_completed=True)
        provider = request.GET.get('provider')
        if provider:
            orders = orders.filter(provider=CardboardProvider.objects.get(name=provider))
        providers = CardboardProvider.objects.all()
        quantities = OrderItemQuantity.objects.all()
        return render(request, 'warehousemanager-all-orders-details.html', locals())


class NewOrder(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_order'

    def get(self, request):
        providers = CardboardProvider.objects.all()
        form = NewOrderForm()

        # deleting orders without items
        orders = Order.objects.all()
        for order in orders:
            if not OrderItem.objects.all().filter(order=order):
                order.delete()

        not_completed_orders = Order.objects.all()

        for n_c_order in not_completed_orders:
            if not n_c_order.is_completed:
                messages.info(request, 'Masz niezakończone zamówienia')

        return render(request, 'warehousemanager-new-order.html', locals())

    def post(self, request):
        form = NewOrderForm(request.POST)
        if form.is_valid():
            provider = form.cleaned_data['provider']
            order_provider_number = form.cleaned_data['order_provider_number']
            date_of_order = form.cleaned_data['date_of_order']

            provider_object = CardboardProvider.objects.get(name=provider)

            new_order = Order.objects.create(provider=provider_object, order_provider_number=int(order_provider_number),
                                             date_of_order=date_of_order)

            new_order.save()

            orders = Order.objects.all()
            if orders:
                order_id = orders.order_by('id').reverse()[0].id
            else:
                order_id = 1
            return redirect('/add-items/{}'.format(order_id))


class DeleteOrder(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.delete_order'

    def get(self, request):
        order_id = int(request.GET.get('order_id'))
        Order.objects.get(id=order_id).delete()

        return redirect('uncompleted-orders')


class NewOrderAdd(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_order'

    def get(self, request):
        provider_num = request.GET.get('provider_num')
        provider = CardboardProvider.objects.get(id=int(provider_num))
        all_orders = Order.objects.all().filter(provider=provider).order_by('order_provider_number')
        num = all_orders.reverse()[0].order_provider_number + 1
        new_order = Order.objects.create(provider=provider, order_provider_number=num,
                                         date_of_order=datetime.datetime.now())
        new_order.save()

        return HttpResponse('')


class NewItemAdd(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_orderitem'

    def get(self, request, order_id):
        form = NewOrderItemForm()
        order = Order.objects.get(id=order_id)
        items = OrderItem.objects.all().filter(order=order)
        last_items = OrderItem.objects.all().reverse()[:5]
        all_items = OrderItem.objects.all()

        if order.is_completed:
            return HttpResponse('Zamówienie zostało już skompletowane')

        return render(request, 'add-item.html', locals())

    def post(self, request, order_id):
        form = NewOrderItemForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/add-items/{}'.format(order_id))
        else:
            return HttpResponse(form.errors)


class OrderItemDelete(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_orderitem'

    def get(self, request, order_id, item_id):
        item_object = OrderItem.objects.get(id=int(item_id))
        item_object.delete()

        return redirect('/add-items/{}'.format(order_id))


class NextOrderNumber(View):
    def get(self, request):
        provider_num = request.GET.get('provider_num')
        provider = CardboardProvider.objects.get(id=int(provider_num))
        all_orders = Order.objects.all().filter(provider=provider).order_by('order_provider_number').reverse()
        if all_orders:
            num = all_orders[0].order_provider_number + 1
        else:
            num = 1

        z = json.dumps(num)

        return HttpResponse(z)


class ProviderForm(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_cardboardprovider'

    def get(self, request):
        form = CardboardProviderForm()
        return render(request, 'new_provider.html', locals())

    def post(self, request):
        form = CardboardProviderForm(request.POST)
        if form.is_valid():
            name = form.cleaned_data['name']
            CardboardProvider.objects.create(name=name)

            return redirect('manage')


class NextItemNumber(View):
    def get(self, request):
        order_num = request.GET.get('order_num')
        order = Order.objects.get(id=int(order_num))
        all_items = OrderItem.objects.all().filter(order=order)

        num = len(all_items)

        return HttpResponse(json.dumps(num + 1))


class CompleteOrder(View):
    def get(self, request):
        order_id = int(request.GET.get('order_id'))
        state = request.GET.get('state')
        order = Order.objects.get(id=order_id)
        order.is_completed = True if state == 'c' else False
        order.save()

        return redirect('all-orders-details')


class UncompletedOrders(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_order'

    def get(self, request):
        orders = Order.objects.filter(is_completed=False)
        providers = CardboardProvider.objects.all()
        return render(request, 'warehousemanager-all-orders-details.html', locals())


class GetItemDetails(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_orderitem'

    def get(self, request):
        item_id = int(request.GET.get('item_id'))
        item = OrderItem.objects.get(id=item_id)

        if item.buyer.all():
            buyer = item.buyer.all()[0].name
        else:
            buyer = ''

        data = {
            'height': item.format_height,
            'width': item.format_width,
            'dimension_one': item.dimension_one,
            'dimension_two': item.dimension_two,
            'dimension_three': item.dimension_three,
            'sort': item.sort,
            'buyer': buyer,
            'weight': item.cardboard_weight,
            'cardboard_type': item.cardboard_type,
            'name': item.name,
            'scores': item.scores
        }

        return HttpResponse(json.dumps(data))


class PrintTest(View):
    def get(self, request):
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer)
        p.drawString(100, 100, 'Hello World.')
        p.showPage()
        p.save()
        lpr = subprocess.Popen("/usr/bin/lpr", stdin=subprocess.PIPE)
        lpr.stdin.write(b'ok')
        buffer.seek(0)

        return FileResponse(buffer, as_attachment=True, filename='hello.pdf')


class OpenFile(View):
    def get(self, request, order_item_id):
        order_item = OrderItem.objects.get(id=order_item_id)

        os.chdir('/home/damian/PycharmProjects/PakerProject/media/paker')

        src = 'zp.docx'
        dst = 'zp{}'.format(order_item.id)

        os.rename(src, 'oko.docx')

        '''document = docx.Document('zp{}'.format(order_item.id))
        print(len(document.sections))

        text = ''

        for p in document.paragraphs:
            if p.text == 'TERMIN REALIZACJI':
                p.text = 'nowy termin'
            text += ' $'
            text += '<br />'
            text += p.text

        text += '<br />'
        text += 'Tabele <br />'

        for t in document.tables:
            text2 = ''
            for i, row in enumerate(t.rows):
                for c in row.cells:
                    if c.text == 'TYP MASZYNY':
                        if order_item.sort in ('201', '202', '203'):
                            c.text = 'f. ' + order_item.sort

                    text2 += c.text

            text += text2

        document.save('zp.docx')'''

        return HttpResponse('')


class NewAllOrders(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_order'

    def get(self, request):
        orders = Order.objects.all()
        provider = request.GET.get('provider')
        if provider:
            orders = orders.filter(provider=CardboardProvider.objects.get(name=provider))
        paginator = Paginator(orders, 3)
        page_number = request.GET.get('page')
        print(page_number)
        page_obj = paginator.get_page(page_number)
        providers = CardboardProvider.objects.all()
        quantities = OrderItemQuantity.objects.all()
        return render(request, 'new-all-orders.html', locals())


class StartPage(View):
    def get(self, request):
        user = request.user
        return render(request, 'start-page.html', locals())

    def post(self, request):
        name = request.POST.get('login')
        password = request.POST.get('password')

        user = authenticate(username=name, password=password)

        if user is not None:
            login(request, user)

        return redirect('start-page')


class LogoutView(View):
    def get(self, request):
        logout(request)

        return redirect('start-page')


class ManageView(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request):
        # title = 'MANAGEMENT'
        # return render(request, 'warehousemanager-manage.html', locals())
        return redirect('punches')


# wszyscy dostawcy
class AllProvidersView(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request):
        providers = CardboardProvider.objects.all()
        return render(request, 'warehousemanager-all-providers.html', locals())


# przelicznik formatów
class FormatConverter(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request):
        return render(request, 'warehousemanager-format-converter.html', locals())


# zarządzanie dostawami
class DeliveriesManagement(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request):
        title = 'DELIVERIES'
        all_deliveries = Delivery.objects.all()
        return render(request, 'warehousemanager-all-deliveries.html', locals())


# szczegóły dostawy
class DeliveryDetails(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request, delivery_id):
        delivery = Delivery.objects.get(id=delivery_id)
        quantities = OrderItemQuantity.objects.filter(delivery=delivery)

        order_item_q_form = OrderItemQuantityForm(initial={'delivery': delivery}, provider=delivery.provider)

        return render(request, 'warehousemanager-delivery-details.html', locals())

    def post(self, request, delivery_id):
        delivery = Delivery.objects.get(id=delivery_id)
        order_item_q_form = OrderItemQuantityForm(delivery.provider, request.POST)

        if order_item_q_form.is_valid():
            delivery = request.POST.get('delivery')
            order_item = request.POST.get('order_item')
            quantity = request.POST.get('quantity')

            new_oiq = OrderItemQuantity.objects.create(delivery=Delivery.objects.get(id=int(delivery)),
                                                       order_item=OrderItem.objects.get(id=int(order_item)),
                                                       quantity=int(quantity))

            new_oiq.save()

            return redirect('/delivery/{}'.format(delivery_id))


# dodawanie dostawy
class DeliveryAdd(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_delivery'

    def get(self, request):
        delivery_form = DeliveryForm()

        return render(request, 'warehousemanager-delivery-add.html', locals())

    def post(self, request):
        delivery_form = DeliveryForm(request.POST)
        if delivery_form.is_valid():
            provider = delivery_form.cleaned_data['provider']
            date_of_delivery = delivery_form.cleaned_data['date_of_delivery']

            new_delivery = Delivery.objects.create(provider=provider, date_of_delivery=date_of_delivery)

            new_delivery.save()

            return redirect('/delivery/{}'.format(new_delivery.id))

        else:
            return HttpResponse('fail')


# dodawanie notatek
class NoteAdd(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_note'

    def get(self, request):
        note_form = NoteForm()

        return render(request, 'warehousemanager-add-note.html', locals())

    def post(self, request):
        note_form = NoteForm(request.POST)
        if note_form.is_valid():
            genre = note_form.cleaned_data['genre']
            title = note_form.cleaned_data['title']
            content = note_form.cleaned_data['content']
            new_note = Note.objects.create(genre=genre, title=title, content=content)

            new_note.save()

            return redirect('notes')


# wszystkie notatki
class AllNotes(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_note'

    def get(self, request):
        all_notes = Note.objects.all()
        return render(request, 'warehousemanager-all-notes.html', locals())


class AbsencesList(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request):
        def month_days_function(month_and_year):
            if month_and_year.month in (1, 3, 5, 7, 8, 10, 12):
                days_num = 31
            elif month_and_year.month == 2:
                if month_and_year.year % 4 == 0:
                    days_num = 29
                else:
                    days_num = 28
            else:
                days_num = 30

            return days_num

        def weekday_of_month_days(some_date, day):
            weekday_ = datetime.datetime.strptime(f'{day}-{some_date.month}-{some_date.year}', '%d-%m-%Y')
            weekday_num = weekday_.weekday()

            return weekday_num

        def today_month():
            today = datetime.datetime.now()

            ab = datetime.datetime.strftime(today, '%d-%m-%Y')
            a_dt = datetime.datetime.strptime(ab, '%d-%m-%Y')

            am = months[int(a_dt.month - 1)]
            ay = a_dt.year

            aaa = am + ' ' + str(ay)

            return aaa

        def previous_and_next_month(some_date_str_yyyy_mm_dd):
            this_month = datetime.datetime.strptime(some_date_str_yyyy_mm_dd, '%Y-%m-%d').month
            this_year = datetime.datetime.strptime(some_date_str_yyyy_mm_dd, '%Y-%m-%d').year

            if int(this_month) == 12:
                next_month_num = 1
                prev_month_num = 11
                next_year = this_year + 1
                prev_year = this_year
            elif int(this_month) == 1:
                next_month_num = 2
                prev_month_num = 12
                next_year = this_year
                prev_year = this_year - 1
            else:
                next_month_num = int(this_month) + 1
                prev_month_num = int(this_month) - 1
                prev_year = this_year
                next_year = this_year
            return months[prev_month_num - 1] + f' {prev_year}', months[next_month_num - 1] + f' {next_year}'

        months = (
            'January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
            'November',
            'December')

        month = request.GET.get('month')

        if not month:
            month = datetime.date.today()
            aa = today_month()
            month_date = datetime.datetime.today()
            day_num = month_date.day
            prev_month, next_month = previous_and_next_month(datetime.date.strftime(month, '%Y-%m-%d'))
        else:
            aa = month
            month_split = month.split()
            new_date = f'01-{months.index(month_split[0]) + 1}-{month_split[1]}'
            month_date = datetime.datetime.strptime(new_date, '%d-%m-%Y')
            if month_date.month == datetime.datetime.today().month and month_date.year == datetime.datetime.today().year:
                day_num = datetime.datetime.today().day
            else:
                day_num = 32
            prev_month, next_month = previous_and_next_month(f'{month_split[1]}-{months.index(month_split[0]) + 1}-01')

        month_year = aa.split()
        str_date = f'{month_year[1]}-{months.index(month_year[0]) + 1}-1'
        month_days = month_days_function(datetime.datetime.strptime(str_date, '%Y-%m-%d'))

        workers = Person.objects.all().filter(
            job_start__lte=datetime.date(int(month_year[1]), months.index(month_year[0]) + 1, month_days))
        workers = workers.exclude(job_end__lte=datetime.date(int(month_year[1]), months.index(month_year[0]) + 1, 1))
        '''to_delete = []
        for worker in workers:
            if worker.job_end:
                if worker.job_end < datetime.date(int(month_year[1]), months.index(month_year[0]) + 1, 1):
                    to_delete.append(worker.id)
                    print(worker.first_name, worker.last_name)
        workers.filter(id__in=to_delete).delete()'''
        absences = Absence.objects.all()

        a = datetime.datetime.strftime(month_date, '%d-%m-%Y')

        month_days = [(x, weekday_of_month_days(datetime.datetime.strptime(a, '%d-%m-%Y'), x)) for x in
                      range(1, month_days_function(month_date) + 1)]

        def month_list(start_date, end_date):
            start_date_str = datetime.datetime.strptime(start_date, '%d-%m-%Y')
            end_date_str = datetime.datetime.strptime(end_date, '%d-%m-%Y')

            month_year = str(start_date_str.month) + ' ' + str(start_date_str.year)
            month_year_end = str(end_date_str.month) + ' ' + str(end_date_str.year)

            months_list = []

            while month_year != month_year_end:
                month_num_arg = datetime.datetime.strftime(start_date_str, '%d-%m-%Y')
                month_num = datetime.datetime.strptime(month_num_arg, '%d-%m-%Y').month
                year = datetime.datetime.strptime(month_num_arg, '%d-%m-%Y').year
                month_name = months[int(month_num) - 1]
                months_list.append(month_name + ' ' + str(year))
                if start_date_str.month in (1, 3, 5, 7, 8, 10, 12):
                    days_num = 31
                elif start_date_str.month == 2:
                    if start_date_str.year % 4 == 0:
                        days_num = 29
                    else:
                        days_num = 28
                else:
                    days_num = 30
                start_date_str = start_date_str + datetime.timedelta(days=days_num)
                month_year = str(start_date_str.month) + ' ' + str(start_date_str.year)

            return months_list

        end_list_condition = next_month == f'{months[datetime.date.today().month]} {str(datetime.date.today().year)}'
        end_plus_31 = datetime.datetime.today() + datetime.timedelta(days=31)
        z = month_list('01-01-2017', datetime.datetime.strftime(end_plus_31, '%d-%m-%Y'))

        return render(request, 'warehousemanager-absenceslist.html', locals())


class AbsencesAndHolidays(View):

    def get(self, request):
        def days_without_work(worker, mm_yyyy_str):
            start = worker.job_start
            end = worker.job_end
            month__ = datetime.datetime.strptime(mm_yyyy_str, '%m-%Y')

            if month__.month in (1, 3, 5, 7, 8, 10, 12):
                days = 31
            elif month__.month == 2:
                days = 28
                if month__.year % 4 == 0:
                    days = 29
            else:
                days = 30

            result_days = []

            if start.month == month__.month and start.year == month__.year:
                result_days = [x for x in range(1, start.day)]
            if end:
                if end.month == month__.month and end.year == month__.year:
                    for y in range(end.day, days + 1):
                        result_days.append(y)

            return worker.id, result_days

        months = (
            'January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
            'November',
            'December')
        month = request.GET.get('month')
        div_date = month.split()
        year = div_date[1]
        month_ = months.index(div_date[0]) + 1
        if month_ != 12:
            absences_objects = Absence.objects.all().filter(absence_date__gte=datetime.date(int(year), month_, 1),
                                                            absence_date__lt=datetime.date(int(year), month_ + 1, 1))
            holiday_objects = Holiday.objects.all().filter(holiday_date__gte=datetime.date(int(year), month_, 1),
                                                           holiday_date__lt=datetime.date(int(year), month_ + 1, 1))
        else:
            absences_objects = Absence.objects.all().filter(absence_date__gte=datetime.date(int(year), month_, 1),
                                                            absence_date__lt=datetime.date(int(year) + 1, 1, 1))
            holiday_objects = Holiday.objects.all().filter(holiday_date__gte=datetime.date(int(year), month_, 1),
                                                           holiday_date__lt=datetime.date(int(year) + 1, 1, 1))

        # deleting vacations on holidays
        for h in holiday_objects:
            absences = Absence.objects.all().filter(absence_date=h.holiday_date)
            if len(absences) > 0:
                for a in absences:
                    a.delete()

        workers = Person.objects.all()
        non_work_days = []

        for w in workers:
            r = days_without_work(w, f'{month_}-{year}')
            if len(r[1]) > 0:
                non_work_days.append(r)

        extra_hours = []
        month_start_date = datetime.date(int(year), month_, 1)
        month_end_date = datetime.date(int(year), month_, 28)
        e_h = ExtraHour.objects.all().filter(extras_date__gte=month_start_date).filter(extras_date__lte=month_end_date)

        for e in e_h:
            extra_hours.append((e.worker.id, e.extras_date.day, float(e.quantity)))

        absences_and_holidays = []
        for a in absences_objects:
            absences_and_holidays.append((a.worker.id, a.absence_date.day, a.absence_type))
        for h in holiday_objects:
            absences_and_holidays.append((-1, h.holiday_date.day, h.name))
        return HttpResponse(json.dumps((absences_and_holidays, non_work_days, extra_hours)))


class GetLocalVar(View):
    def get(self, request, variable_name):
        if os.environ[variable_name]:
            return HttpResponse(json.dumps(os.environ[variable_name]))
        else:
            return redirect('manage')


class AbsenceAdd(LoginRequiredMixin, View):
    login_url = '/'

    def get(self, request):
        short_absence_form = AbsenceForm
        workers = Person.objects.all()
        reasons = ABSENCE_TYPES
        return render(request, 'warehousemanager-add-absence.html', locals())

    def post(self, request):
        short_absence_form = AbsenceForm(request.POST)
        if short_absence_form.is_valid():
            worker = short_absence_form.cleaned_data['worker']
            absence_date = short_absence_form.cleaned_data['absence_date']
            absence_type = short_absence_form.cleaned_data['absence_type']

            new_absence = Absence.objects.create(worker=worker, absence_date=absence_date, absence_type=absence_type)

            new_absence.save()

            return redirect('absence-list')

        else:
            worker = request.POST.get('worker')
            first_day = request.POST.get('first_day')
            last_day = request.POST.get('last_day')
            absence_type = request.POST.get('type')

            worker_s = worker.split()
            worker_object = Person.objects.filter(first_name=worker_s[0], last_name=worker_s[1])[0]
            first_day_date = datetime.datetime.strptime(first_day, '%Y-%m-%d')
            last_day_date = datetime.datetime.strptime(last_day, '%Y-%m-%d')

            safety_counter = 0
            while first_day_date != last_day_date:
                if safety_counter < 15:
                    safety_counter += 1
                    if first_day_date.weekday() < 5:
                        new_absence = Absence(worker=worker_object, absence_date=first_day_date,
                                              absence_type=absence_type)
                        new_absence.save()
                    else:
                        if absence_type == 'UZ':
                            new_absence = Absence(worker=worker_object, absence_date=first_day_date,
                                                  absence_type=absence_type)
                            new_absence.save()
                        else:
                            pass
                    first_day_date = first_day_date + datetime.timedelta(days=1)
                    if first_day_date == last_day_date:
                        new_absence = Absence(worker=worker_object, absence_date=first_day_date,
                                              absence_type=absence_type)
                        new_absence.save()
                else:
                    break

            return redirect('absence-list')


class PunchesList(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punch'

    def get(self, request):
        punches = Punch.objects.all().order_by('type', 'type_letter', 'type_num')
        punch_types = PUNCH_TYPES
        title = 'PUNCHES'

        return render(request, 'warehousemanager-punches-list.html', locals())


class PunchAdd(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punch'

    def get(self, request):
        punch_form = PunchForm()

        return render(request, 'warehousemanager-punch-add.html', locals())

    def post(self, request):
        punch_form = PunchForm(request.POST)

        if punch_form.is_valid():
            punch_type = punch_form.cleaned_data['type']
            type_letter = type_num = punch_form.cleaned_data['type_letter']
            type_num = punch_form.cleaned_data['type_num']
            name = punch_form.cleaned_data['name']
            dimension_one = punch_form.cleaned_data['dimension_one']
            dimension_two = punch_form.cleaned_data['dimension_two']
            dimension_three = punch_form.cleaned_data['dimension_three']
            quantity = punch_form.cleaned_data['quantity']
            size_one = punch_form.cleaned_data['size_one']
            size_two = punch_form.cleaned_data['size_two']
            cardboard = punch_form.cleaned_data['cardboard']
            pressure_large = punch_form.cleaned_data['pressure_large']
            pressure_small = punch_form.cleaned_data['pressure_small']
            wave_direction = punch_form.cleaned_data['wave_direction']
            customers = punch_form.cleaned_data['customers']

            new_punch = Punch.objects.create(type=punch_type, type_num=type_num, dimension_one=dimension_one,
                                             dimension_two=dimension_two, dimension_three=dimension_three,
                                             quantity=quantity, size_one=size_one, size_two=size_two,
                                             cardboard=cardboard, pressure_large=pressure_large,
                                             pressure_small=pressure_small, wave_direction=wave_direction, name=name,
                                             type_letter=type_letter)

            new_punch.save()

            return redirect('punches')


class PunchDetails(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punch'

    def get(self, request, punch_id):
        p = get_object_or_404(Punch, id=punch_id)
        production = PunchProduction.objects.filter(punch=p)

        wear = 0
        for pr in production:
            wear += pr.quantity

        return render(request, 'warehousemanager-punch-details.html', locals())


class PunchEdit(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punch'

    def get(self, request, punch_id):
        p = get_object_or_404(Punch, id=punch_id)
        punch_form = PunchForm(instance=p)
        edit = True

        return render(request, 'warehousemanager-punch-add.html', locals())

    def post(self, request, punch_id):
        punch_form = PunchForm(request.POST)

        if punch_form.is_valid():
            punch_type = punch_form.cleaned_data['type']
            type_letter = type_num = punch_form.cleaned_data['type_letter']
            type_num = punch_form.cleaned_data['type_num']
            name = punch_form.cleaned_data['name']
            dimension_one = punch_form.cleaned_data['dimension_one']
            dimension_two = punch_form.cleaned_data['dimension_two']
            dimension_three = punch_form.cleaned_data['dimension_three']
            quantity = punch_form.cleaned_data['quantity']
            size_one = punch_form.cleaned_data['size_one']
            size_two = punch_form.cleaned_data['size_two']
            cardboard = punch_form.cleaned_data['cardboard']
            pressure_large = punch_form.cleaned_data['pressure_large']
            pressure_small = punch_form.cleaned_data['pressure_small']
            wave_direction = punch_form.cleaned_data['wave_direction']
            customers = punch_form.cleaned_data['customers']

            edited_punch = Punch.objects.get(id=punch_id)

            edited_punch.type = punch_type
            edited_punch.type_num = type_num
            edited_punch.dimension_one = dimension_one
            edited_punch.dimension_two = dimension_two
            edited_punch.dimension_three = dimension_three
            edited_punch.quantity = quantity
            edited_punch.size_one = size_one
            edited_punch.size_two = size_two
            edited_punch.cardboard = cardboard
            edited_punch.pressure_small = pressure_small
            edited_punch.pressure_large = pressure_large
            edited_punch.wave_direction = wave_direction
            edited_punch.name = name
            edited_punch.type_letter = type_letter

            edited_punch.save()

            print(customers)

            return redirect('punches')


class PunchDelete(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.delete_punch'

    def get(self, request, punch_id):
        p = get_object_or_404(Punch, id=punch_id)
        production = PunchProduction.objects.filter(punch=p)
        if production:
            return redirect('announcement')
        p.delete()

        return redirect('punches')


class AddBuyer(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_buyer'

    def get(self, request):
        buyer_form = BuyerForm()

        return render(request, 'warehousemanager-add-buyer.html', locals())

    def post(self, request):
        buyer_form = BuyerForm(request.POST)

        if buyer_form.is_valid():
            name = buyer_form.cleaned_data['name']
            new_buyer = Buyer.objects.create(name=name)
            new_buyer.save()

            return redirect('manage')


class BuyersList(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_buyer'

    def get(self, request):
        buyers = Buyer.objects.all()

        return render(request, 'warehousemanager-buyers-list.html', locals())


class PunchProductions(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punchproduction'

    def get(self, request):
        production = PunchProduction.objects.all()

        return render(request, 'warehousemanager-punch-production.html', locals())


class PunchProductionAdd(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punchproduction'

    def get(self, request):
        punch_id = request.GET.get('punch_id')
        if punch_id:
            punch = Punch.objects.get(id=int(punch_id))
            form = PunchProductionForm(initial={'punch': punch})

        else:
            form = PunchProductionForm()

        return render(request, 'warehousemanager-punch-production-add.html', locals())

    def post(self, request):
        form = PunchProductionForm(request.POST)

        if form.is_valid():
            punch = form.cleaned_data['punch']
            worker = form.cleaned_data['worker']
            date_start = form.cleaned_data['date_start']
            date_end = form.cleaned_data['date_end']
            quantity = form.cleaned_data['quantity']
            comments = form.cleaned_data['comments']

            production = PunchProduction.objects.create(punch=punch, worker=worker, date_start=date_start,
                                                        date_end=date_end, quantity=quantity, comments=comments)
            production.save()

            return redirect('punches')


class CardboardUsed(View):
    def get(self, request, cardboard_stock_id):
        cardboard = OrderItemQuantity.objects.get(id=cardboard_stock_id)
        if cardboard.is_used:
            return HttpResponse(json.dumps(True))
        else:
            return HttpResponse(json.dumps(False))


class StockManagement(PermissionRequiredMixin, View):
    permission_required = 'warehousemanager.view_punchproduction'

    def get(self, request):
        stocks = OrderItemQuantity.objects.filter(is_used=False)
        history_stocks = OrderItemQuantity.objects.filter(is_used=True)

        return render(request, 'warehousemanager-stock-management.html', locals())


class Announcement(View):
    def get(self, request):
        return render(request, 'warehousemanager-announcement.html')


class ChangeOrderState(View):
    def get(self, request):
        order_item_id = request.GET.get('order_item_id')
        order_item = OrderItem.objects.get(id=int(order_item_id))
        print(order_item)
        if order_item.is_completed:
            order_item.is_completed = False
        else:
            order_item.is_completed = True

        order_item.save()

        return HttpResponse(json.dumps(order_item.is_completed))


class ProductionView(View):
    def get(self, request):
        items_to_do = OrderItem.objects.filter(is_completed=True)
        return render(request, 'warehousemanager-production-status.html', locals())


class OrderItemDetails(View):
    def get(self, request, order_item_id):
        order_item = OrderItem.objects.get(id=order_item_id)
        return render(request, 'warehousemanager-order-item-details.html', locals())

